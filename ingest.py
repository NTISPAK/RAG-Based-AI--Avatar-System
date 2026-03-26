"""
Document Ingestion Script for NTIS Policy RAG Chatbot
Loads DOCX documents, chunks them, creates embeddings, and stores in Qdrant.

Primary sources:
  - Updated Study Visa Team Document.docx
  - NIRA Updated Document.docx
"""
import os
from dotenv import load_dotenv
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams

load_dotenv()

# Configuration — primary document sources
BASE_DIR = os.path.dirname(__file__)
DOC_PATHS = [
    os.path.join(BASE_DIR, "Updated Study Visa Team Document .docx"),
    os.path.join(BASE_DIR, "NIRA Updated Document .docx"),
]
QDRANT_URL = os.getenv("QDRANT_URL", "http://localhost:6333")
COLLECTION_NAME = "policy_docs"


def load_docx(path: str) -> list[Document]:
    """Load a .docx file and return LangChain Document objects (one per section)."""
    doc = DocxDocument(path)
    full_text = []
    current_section = ""
    current_heading = os.path.basename(path)

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            if current_section.strip():
                full_text.append(Document(
                    page_content=current_section.strip(),
                    metadata={"source": os.path.basename(path), "heading": current_heading}
                ))
            current_heading = para.text or current_heading
            current_section = para.text + "\n"
        else:
            current_section += para.text + "\n"

    if current_section.strip():
        full_text.append(Document(
            page_content=current_section.strip(),
            metadata={"source": os.path.basename(path), "heading": current_heading}
        ))

    return full_text


def chunk_documents(documents: list[Document], chunk_size=800, chunk_overlap=150) -> list[Document]:
    """Split documents into smaller chunks with overlap (no external dependency)."""
    separators = ["\n\n", "\n", ". ", " "]
    chunks = []

    for doc in documents:
        text = doc.page_content
        if len(text) <= chunk_size:
            chunks.append(doc)
            continue

        parts = _recursive_split(text, separators, chunk_size)
        # Add overlap between consecutive parts
        for i, part in enumerate(parts):
            if i > 0 and chunk_overlap > 0:
                prev_tail = parts[i - 1][-chunk_overlap:]
                part = prev_tail + part
            if part.strip():
                chunks.append(Document(
                    page_content=part.strip(),
                    metadata=doc.metadata.copy()
                ))

    return chunks


def _recursive_split(text: str, separators: list[str], chunk_size: int) -> list[str]:
    """Recursively split text using separators until chunks are below chunk_size."""
    if len(text) <= chunk_size:
        return [text]

    sep = separators[0] if separators else ""
    remaining_seps = separators[1:] if len(separators) > 1 else []

    if sep:
        pieces = text.split(sep)
    else:
        # Last resort: hard split by char count
        return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

    result = []
    current = ""
    for piece in pieces:
        candidate = (current + sep + piece) if current else piece
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            if current:
                result.append(current)
            if len(piece) > chunk_size and remaining_seps:
                result.extend(_recursive_split(piece, remaining_seps, chunk_size))
            else:
                current = piece
                continue
            current = ""
    if current:
        result.append(current)

    return result


def main():
    print("=" * 60)
    print("NTIS Policy Document Ingestion (Updated)")
    print("=" * 60)

    # 1. Load documents
    all_documents = []
    for doc_path in DOC_PATHS:
        print(f"\n[1] Loading: {os.path.basename(doc_path)}")
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"Document not found: {doc_path}")
        docs = load_docx(doc_path)
        print(f"    → {len(docs)} sections loaded")
        all_documents.extend(docs)

    print(f"\n    Total sections across all documents: {len(all_documents)}")

    # 2. Split into chunks
    print("\n[2] Splitting into chunks...")
    chunks = chunk_documents(all_documents, chunk_size=800, chunk_overlap=150)
    print(f"    Created {len(chunks)} chunks")

    # 3. Initialize embeddings model
    print("\n[3] Loading embeddings model...")
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': True}
    )
    print("    Model loaded: all-MiniLM-L6-v2")

    # 4. Flush old collection and create fresh one
    print(f"\n[4] Flushing Qdrant at {QDRANT_URL}...")
    client = QdrantClient(url=QDRANT_URL)

    existing = client.get_collections().collections
    for col in existing:
        client.delete_collection(collection_name=col.name)
        print(f"    Deleted old collection: {col.name}")

    client.create_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=VectorParams(size=384, distance=Distance.COSINE)
    )
    print(f"    Created fresh collection: {COLLECTION_NAME}")

    # 5. Store vectors in Qdrant
    print("\n[5] Storing vectors in Qdrant...")
    vector_store = QdrantVectorStore.from_documents(
        documents=chunks,
        embedding=embeddings,
        url=QDRANT_URL,
        collection_name=COLLECTION_NAME,
    )
    print(f"    Stored {len(chunks)} vectors")

    # 6. Verify
    info = client.get_collection(collection_name=COLLECTION_NAME)
    print(f"\n    Verification — points in collection: {info.points_count}")

    print("\n" + "=" * 60)
    print("SUCCESS! Document ingestion complete.")
    print(f"Collection : {COLLECTION_NAME}")
    print(f"Documents  : {', '.join(os.path.basename(p) for p in DOC_PATHS)}")
    print(f"Vectors    : {len(chunks)}")
    print("=" * 60)


if __name__ == "__main__":
    main()
